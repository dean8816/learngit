import docx

fn = r'D:\test.docx'
doc = docx.Document(fn)
print ("Hello Python")

# 按段落读取全部数据
# for paragraph in doc.paragraphs:
#     print("段落全部内容",paragraph.text)

# print('Hello World')

# 按表格读取全部数据
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            print(cell.text)

table_num = len(doc.tables)
# 获取文档的表格个数
print(table_num)

a = 1000

# table_0 = doc.tables[0]
# # 选取第一个表
# table_rows = len(table_0.rows)
# # 获取第一个表的行数
# print(table_rows)

# tab = doc.tables[0].rows[0].cells[0]
# # 获取第一张表第一行第一列数据
# print(tab.text)

# par = doc.paragraphs[2]
# # 读取第三段数据
# print(par.text)